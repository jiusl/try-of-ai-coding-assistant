# -*- coding: utf-8 -*-
"""
生成 "Try — AI编程助手" 项目实习报告 (.docx)
参考模板：西北农林科技大学项目实习报告（参考模板）.docx
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)  # 小四号
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)  # 2个字符

# 设置页面边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_title_center(text, font_name='黑体', font_size=Pt(22), bold=False, space_after=Pt(6)):
    """添加居中标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def add_body(text, font_name='宋体', font_size=Pt(12), bold=False, first_indent=True):
    """添加正文段落 (小四号宋体, 1.5倍行距, 首行缩进2字符)"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    else:
        p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def add_chapter_title(text, font_name='黑体', font_size=Pt(16)):
    """添加章标题 (三号黑体加粗居中)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def add_section_title(text, font_name='黑体', font_size=Pt(14)):
    """添加节标题 (四号黑体居左)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def add_subsection_title(text, font_name='宋体', font_size=Pt(12)):
    """添加次节标题 (小四号宋体加粗居左)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def add_table_caption(text):
    """添加表题 (五号宋体居中)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_figure_caption(text):
    """添加图题 (五号宋体居中)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def new_page():
    """插入分页符"""
    doc.add_page_break()

def add_code_block(code_text):
    """添加代码块 (五号Consolas/宋体)"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(1)
    # 添加灰色底纹
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ============================================================
# 封面
# ============================================================
# 空行
for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)

# 标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run('本科实训项目设计（论文）报告')
run.font.name = '黑体'
run.font.size = Pt(22)  # 二号
run.font.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 副标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.space_before = Pt(18)
run = p.add_run('基于Effect-TS架构的AI编程助手系统设计与实现')
run.font.name = '黑体'
run.font.size = Pt(18)
run.font.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 空行
for _ in range(3):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)

# 信息表
info_data = [
    ('学    号：', '2023013314'),
    ('姓    名：', '刘聍徽'),
    ('学    院：', '信息工程学院'),
    ('专    业：', '计算机科学与技术'),
    ('指导教师：', '___________'),
    ('起止日期：', f'2026年6月15日 - 2026年6月30日'),
]

for label, value in info_data:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(label + value)
    run.font.name = '宋体'
    run.font.size = Pt(14)  # 四号
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

new_page()

# ============================================================
# 中文摘要
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.line_spacing = 1.5
run = p.add_run('摘    要')
run.font.name = '黑体'
run.font.size = Pt(16)  # 三号
run.font.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 空一行
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)

abstract_text = (
    "随着人工智能技术的快速发展，大语言模型（LLM）在软件开发领域的应用日益深入。"
    "GitHub Copilot、Cursor等商业AI编程助手虽然功能强大，但存在闭源、数据隐私、"
    "定制化受限等问题。本文设计并实现了一款开源、本地化、可扩展的AI编程助手系统——Try，"
    "旨在为开发者提供自主可控的智能编程辅助工具。"
    "系统基于Bun运行时和TypeScript语言开发，采用Effect-TS函数式框架构建核心架构，"
    "实现了依赖注入、类型安全错误处理和响应式流式数据管道。"
    "系统包含八大核心模块：Agent智能体系统支持8种专用角色（Chat/Builder/Coder/Reviewer/"
    "Tester/Refactor/Researcher/Orchestrator），实现任务分解与委托执行；"
    "工具系统提供14个内置Python工具（文件读写、代码搜索、Shell命令、网页抓取等），"
    "采用JSON标准输入输出协议实现跨进程通信；Skill知识注入系统支持动态加载Markdown定义的专业知识模块；"
    "记忆系统通过自动压缩和嵌入检索维护跨会话上下文；会话管理系统基于SQLite数据库（含FTS5全文搜索）"
    "实现会话持久化和工作目录(Workspace)隔离；"
    "权限系统采用基于规则的细粒度控制引擎，配合RBAC角色管理和操作确认机制保障安全；"
    "认证授权系统支持密码注册登录和JWT令牌认证；前端采用React 19 + Vite 6 + Chakra UI v3 "
    "构建响应式Web界面，结合SSE流式推送实现工具调用实时可视化；"
    "系统提供CLI命令行和Web UI双模式交互，并支持Docker容器化一键部署。"
    "在Windows环境下特别解决了Python子进程的中文路径编码兼容性问题。"
    "经测试验证，系统在多模型（OpenAI/Anthropic/DeepSeek/Ollama/llama.cpp）支持下"
    "能够完成代码生成、文件编辑、项目重构、网页研究等多种编程辅助任务。"
)
add_body(abstract_text, first_indent=True)

# 空一行
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)

# 关键词
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run('关键词：')
run.font.name = '宋体'
run.font.size = Pt(12)
run.font.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run('AI编程助手，大语言模型，Effect-TS，Agent智能体，工具系统，TypeScript，Bun运行时')
run.font.name = '宋体'
run.font.size = Pt(12)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

new_page()

# ============================================================
# 英文摘要
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.line_spacing = 1.5
run = p.add_run('ABSTRACT')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)  # 三号
run.font.bold = True

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)

abstract_en = (
    "With the rapid advancement of artificial intelligence, Large Language Models (LLMs) "
    "are playing an increasingly important role in software development. While commercial "
    "AI coding assistants such as GitHub Copilot and Cursor offer powerful features, they "
    "suffer from closed-source limitations, data privacy concerns, and restricted customization. "
    "This paper presents the design and implementation of Try, an open-source, locally-deployable, "
    "and extensible AI-powered coding assistant system that empowers developers with autonomous "
    "and controllable intelligent programming support. "
    "Built on the Bun runtime and TypeScript, the system adopts the Effect-TS functional framework "
    "for its core architecture, implementing dependency injection, type-safe error handling, and "
    "reactive streaming data pipelines. The system comprises eight core modules: an Agent system "
    "with 8 specialized roles (Chat/Builder/Coder/Reviewer/Tester/Refactor/Researcher/Orchestrator) "
    "supporting task decomposition and delegation; a tool system providing 14 built-in Python tools "
    "(file I/O, code search, shell commands, web scraping, etc.) using JSON stdin/stdout protocol "
    "for cross-process communication; a Skill system for dynamic knowledge injection via Markdown-defined "
    "expertise modules; a memory system maintaining cross-session context through automatic compression "
    "and embedding retrieval; a session management system based on SQLite (with FTS5 full-text search) "
    "implementing session persistence and workspace isolation; a permission system with rule-based "
    "fine-grained control, RBAC role management, and operation confirmation mechanisms; an authentication "
    "system supporting password registration/login with JWT tokens; a responsive Web UI built with "
    "React 19, Vite 6, and Chakra UI v3 featuring SSE streaming and real-time tool call visualization; "
    "and dual-mode interaction via CLI and Web UI with Docker containerization for one-click deployment. "
    "Special attention was given to resolving Chinese path encoding compatibility issues with Python "
    "subprocesses on Windows. Testing confirms that the system, supporting multiple models "
    "(OpenAI/Anthropic/DeepSeek/Ollama/llama.cpp), effectively handles code generation, file editing, "
    "project refactoring, and web research tasks."
)
add_body(abstract_en, font_name='Times New Roman', first_indent=True)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run('KEY WORDS: ')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.font.bold = True
run = p.add_run('AI Coding Assistant, Large Language Model, Effect-TS, Agent, Tool System, TypeScript, Bun Runtime')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

new_page()

# ============================================================
# 目录页
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run('目    录')
run.font.name = '黑体'
run.font.size = Pt(16)
run.font.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 目录内容（简化）
toc_items = [
    ('摘    要', 'I'),
    ('ABSTRACT', 'II'),
    ('目    录', 'III'),
    ('第一章 绪论', '1'),
    ('  1.1 课题背景和意义', '1'),
    ('  1.2 国内外研究现状', '2'),
    ('  1.3 本文研究内容', '4'),
    ('第二章 系统需求分析', '5'),
    ('  2.1 功能需求', '5'),
    ('  2.2 非功能需求', '6'),
    ('第三章 系统架构设计', '7'),
    ('  3.1 总体架构', '7'),
    ('  3.2 Effect-TS函数式框架', '8'),
    ('  3.3 模块分层设计', '9'),
    ('第四章 Agent智能体系统', '10'),
    ('  4.1 Agent体系设计', '10'),
    ('  4.2 执行引擎', '11'),
    ('  4.3 Agent委托机制', '12'),
    ('第五章 工具系统', '13'),
    ('  5.1 工具架构设计', '13'),
    ('  5.2 Python执行协议', '14'),
    ('  5.3 内置工具实现', '15'),
    ('第六章 Skill与记忆系统', '16'),
    ('  6.1 Skill知识注入', '16'),
    ('  6.2 记忆系统', '17'),
    ('第七章 会话管理与安全体系', '18'),
    ('  7.1 会话管理', '18'),
    ('  7.2 Workspace工作目录', '19'),
    ('  7.3 权限与安全', '19'),
    ('第八章 前端与交互设计', '20'),
    ('  8.1 React前端架构', '20'),
    ('  8.2 CLI命令行界面', '21'),
    ('第九章 部署与运维', '22'),
    ('  9.1 Docker容器化', '22'),
    ('  9.2 构建与发布', '22'),
    ('第十章 测试与验证', '23'),
    ('  10.1 功能测试', '23'),
    ('  10.2 兼容性测试', '24'),
    ('第十一章 总结与展望', '25'),
    ('  11.1 工作总结', '25'),
    ('  11.2 工作展望', '26'),
    ('参考文献', '27'),
    ('致    谢', '28'),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(item)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

new_page()

# ============================================================
# 第一章 绪论
# ============================================================
add_chapter_title('第一章 绪论')

# 1.1
add_section_title('1.1 课题背景和意义')

add_body(
    "在当今数字化时代，软件开发已成为推动社会进步和经济发展的重要引擎。"
    "随着软件系统复杂度的不断提升，开发者面临着日益增长的编码任务量、代码质量保障、"
    "以及持续学习新技术栈等多重挑战。传统的集成开发环境（IDE）虽然提供了代码补全、"
    "语法检查等基础辅助功能，但其智能化程度有限，无法理解开发者的高层意图并给出"
    "有意义的代码建议。"
)

add_body(
    "2022年底，以ChatGPT为代表的大语言模型（Large Language Models, LLMs）的出现，"
    "彻底改变了人机交互的方式。LLMs展现出了强大的代码理解和生成能力，能够在多种编程语言"
    "中完成从函数级到项目级的代码编写任务。GitHub Copilot作为最早将LLM能力集成到IDE中"
    "的产品，其2023年的用户已超过100万，生成代码量占开发者总代码量的46%以上。"
    "Cursor等新一代AI-first IDE进一步模糊了传统编码与AI辅助的边界。这些工具的成功表明，"
    "AI编程助手已从「锦上添花」演变为「不可或缺」的开发基础设施。"
)

add_body(
    "然而，现有的商业AI编程助手存在若干痛点：第一，代码数据需要上传至云端服务器，"
    "对于涉及商业机密、国家安全或隐私保护的软件开发场景存在数据泄露风险；"
    "第二，闭源系统难以根据特定领域需求进行深度定制，开发者无法自主控制模型选择、"
    "工具链配置和安全策略；第三，高昂的订阅费用对个人开发者和学生群体构成了经济负担；"
    "第四，单一模型的AI助手在面对特定领域问题时能力不足，缺乏多模型灵活切换和编排的能力。"
)

add_body(
    "基于以上背景，本课题旨在设计并实现Try——一款开源、本地化、可灵活扩展的AI编程助手系统。"
    "该系统具备以下核心价值：（1）支持本地部署，开发者数据不出本机，保障代码安全与隐私；"
    "（2）采用模块化、可插拔的架构设计，允许用户根据需求选择不同的LLM提供商、定制工具集和安全策略；"
    "（3）提供8种专用Agent角色和Agent间任务委托机制，实现复杂编程任务的智能编排；"
    "（4）开源免费，降低AI辅助编程的使用门槛，为学术研究和教学实践提供开放的实验平台。"
    "本课题的研究对于推动AI在软件开发领域的民主化、可定制化具有重要的实践意义。"
)

# 1.2
add_section_title('1.2 国内外研究现状')

add_subsection_title('1.2.1 商业AI编程助手')
add_body(
    "在商业领域，GitHub Copilot由GitHub与OpenAI合作开发，基于Codex模型（GPT-3的后代），"
    "于2021年6月发布技术预览版，2022年6月正式上线。Copilot支持VS Code、JetBrains等主流IDE，"
    "能够根据注释和上下文生成完整的函数实现。截至2024年，Copilot已集成GPT-4o模型，支持"
    "多文件编辑、代码审查和对话式交互。Cursor编辑器基于VS Code内核，深度整合AI能力，"
    "提供了Inline editing、Chat、Composer等多种交互模式，支持Claude和GPT系列模型。"
    "JetBrains AI Assistant则深度整合在IntelliJ平台中，提供上下文感知的代码建议。"
    "Amazon CodeWhisperer（现Amazon Q Developer）专注于AWS生态的代码辅助。"
)

add_subsection_title('1.2.2 开源AI编程助手')
add_body(
    "在开源领域，OpenCode项目是一个重要的参考实现。OpenCode由开源社区维护，采用Effect-TS框架"
    "构建，支持多模型后端（OpenAI、Anthropic、Ollama等），提供了Agent/Tool/Skill三大核心抽象。"
    "其架构设计强调类型安全、可组合性和可测试性，为开源AI编程助手树立了设计典范。"
    "Continue.dev是一个面向IDE的开源AI助手插件，支持本地和远程模型，提供可配置的斜杠命令和上下文管理。"
    "Aider是一款基于命令行的AI结对编程工具，专注于Git感知的代码编辑，支持多文件修改和自动提交。"
    "Cody by Sourcegraph结合了代码搜索与LLM能力，能够理解整个代码仓库的上下文。"
    "TabbyML、Ollama等项目则专注于本地模型部署和推理优化，降低了本地LLM推理的门槛。"
)

add_subsection_title('1.2.3 关键技术研究现状')
add_body(
    "在Agent智能体技术方面，ReAct（Reasoning + Acting）模式已成为LLM Agent的主流范式，"
    "通过交替进行推理和行动，使模型能够使用外部工具解决复杂问题。OpenAI的Function Calling、"
    "Anthropic的Tool Use机制为模型调用外部工具提供了标准化的接口。AutoGPT、MetaGPT等"
    "项目探索了多Agent协作完成复杂任务的可行性。在工具系统方面，越来越多的研究关注"
    "如何让LLM安全、高效地使用外部工具。在权限控制方面，沙箱执行、命令审批、文件访问控制"
    "等机制成为保障AI编程助手安全性的必备组件。"
)

# 1.3
add_section_title('1.3 本文研究内容')

add_body(
    "本文围绕开源AI编程助手系统的设计与实现，主要研究内容包括以下几个方面："
)

add_body(
    "（1）系统架构设计：基于Effect-TS函数式框架，设计分层、可组合的系统架构，实现依赖注入、"
    "类型安全错误处理和响应式流式数据管道。将系统拆解为Agent、Tool、Skill、Memory、Session、"
    "Permission、Provider、Server、CLI等独立模块，各模块通过明确的接口进行交互。"
)

add_body(
    "（2）Agent智能体系统：设计多角色Agent体系，实现Chat（通用对话）、Builder（全栈开发）、"
    "Coder（代码编写）、Reviewer（代码审查）、Tester（测试运行）、Refactor（代码重构）、"
    "Researcher（信息检索）、Orchestrator（任务编排）8种专用Agent。研究Agent间的任务委托"
    "与通信机制，设计迭代执行引擎，支持Agent自主规划-执行-反馈循环，并处理最大迭代、"
    "权限拒绝等异常终止场景。"
)

add_body(
    "（3）工具系统：设计可扩展的工具注册与加载机制。将14个核心工具（文件读写、代码搜索、"
    "Shell命令执行、网页抓取、记忆存储与检索等）从TypeScript实现重构为Python实现，建立"
    "JSON标准输入输出跨进程通信协议。解决Windows平台中文路径编码兼容性问题。设计工具"
    "TOOL.md元数据规范，实现工具类型安全校验和参数自动验证。"
)

add_body(
    "（4）前后端与交互设计：实现基于React 19 + Vite 6 + Chakra UI v3的现代化前端界面，"
    "支持多会话管理、Agent动态切换、Workspace可视化选择、工具调用实时显示。基于SSE（Server-Sent Events）"
    "协议实现流式输出推送，确保AI响应的低延迟感知。设计CLI交互式命令行界面，支持丰富的"
    "对话管理命令和终端着色输出。实现用户认证、会话持久化等功能。"
)

add_body(
    "（5）安全与运维：设计基于规则的权限引擎，支持文件模式和命令模式的白名单/黑名单配置，"
    "集成敏感操作确认机制。实现Workspace工作目录隔离，防止Agent越权访问系统文件。"
    "提供Docker容器化部署方案（Dockerfile + docker-compose），编写完整的部署文档。"
    "实现后端TypeScript代码的独立二进制编译（Bun compile），简化分发与部署。"
)

new_page()

# ============================================================
# 第二章 系统需求分析
# ============================================================
add_chapter_title('第二章 系统需求分析')

add_section_title('2.1 功能需求')

add_body(
    "根据AI编程助手的使用场景和目标用户群体，系统需要满足以下核心功能需求："
)

add_body(
    "（1）多模型对话：系统应支持多种LLM提供商，包括OpenAI（gpt-4o系列）、"
    "Anthropic（Claude系列）、DeepSeek（deepseek-v4系列）、Ollama本地部署模型、"
    "以及通过node-llama-cpp加载的本地GGUF格式模型。用户可在运行时动态切换模型，"
    "无需重启系统。系统需支持流式输出，使得用户能够实时看到AI的生成内容。"
)

add_body(
    "（2）Agent智能体：提供8种专用Agent角色，每种Agent拥有独立的系统提示词、"
    "工具集合和能力边界。Agent应能根据用户输入的上下文自适应调整行为策略。"
    "支持Agent间的任务委托，允许Builder和Orchestrator将子任务分配给专业Agent执行。"
    "Agent执行引擎需支持多轮迭代，在达到最大迭代次数时正确终止并反馈状态。"
)

add_body(
    "（3）工具调用：提供文件读写（read_file、write_file、edit_file）、代码搜索"
    "（glob模式匹配、grep内容搜索）、命令执行（run_command、read_command）、"
    "网页抓取（fetch_webpage）、记忆管理（recall、remember）、Skill获取"
    "（list_skills、get_skill）、内部推理（think）等14个内置工具。工具需支持"
    "参数验证、执行超时控制和错误处理。"
)

add_body(
    "（4）会话管理：支持多会话创建、切换和删除，会话元数据自动命名（基于AI生成标题）。"
    "每个会话可独立设置工作目录（Workspace），Agent的所有文件操作限定在该工作目录内。"
    "会话数据（消息历史、配置参数、workspace路径）需持久化到SQLite数据库，支持FTS5全文搜索。"
)

add_body(
    "（5）Web UI界面：提供响应式的浏览器交互界面，包含会话侧边栏、对话面板、"
    "Agent选择器、Workspace工作路径选择器、设置抽屉等组件。支持实时流式显示AI回复，"
    "并以可视化方式展示工具调用过程（工具名称、参数、执行结果）。提供用户注册登录、"
    "模型配置、权限策略配置等功能。"
)

add_section_title('2.2 非功能需求')

add_body(
    "（1）性能需求：SSE流式输出的首字节延迟应在500ms以内。工具执行（如grep搜索）"
    "的超时时间应可配置，默认为30秒。前端页面加载时间（First Contentful Paint）应在2秒以内。"
)

add_body(
    "（2）安全性需求：API Key等敏感配置信息应通过auth.json文件存储并加入.gitignore，"
    "防止泄露到版本控制系统。用户密码使用bcrypt哈希存储。Agent的文件操作需经过权限引擎"
    "检查，敏感操作（删除文件、修改系统配置等）需要用户手动确认。Workspace工作目录机制"
    "应防止Agent越权访问系统文件。JWT认证令牌需设置合理过期时间。"
)

add_body(
    "（3）可扩展性需求：Agent、Tool、Skill三大核心组件均需支持插件化注册。用户可通过"
    "在指定目录下创建TOOL.md和实现脚本即可添加自定义工具。Skill系统支持通过Markdown文件"
    "注入领域知识，无需修改核心代码。系统需提供明确的内置/用户分界目录结构。"
)

add_body(
    "（4）跨平台兼容性需求：系统应在Windows、macOS、Linux三大平台正常运行。"
    "特别需要解决Windows平台上的Python子进程中文路径编码问题（cp936与UTF-8不匹配）。"
    "Python工具脚本需兼容Python 3.10及以上版本。"
)

add_body(
    "（5）运维需求：提供Docker镜像构建和docker-compose一键部署能力。提供独立的"
    "二进制编译能力（Bun compile），方便无Bun运行时的环境使用。提供健康检查端点"
    "和基本的监控指标（Prometheus格式）。支持通过环境变量覆盖配置文件参数。"
)

new_page()

# ============================================================
# 第三章 系统架构设计
# ============================================================
add_chapter_title('第三章 系统架构设计')

add_section_title('3.1 总体架构')

add_body(
    "Try系统采用分层模块化架构，自上而下分为交互层、业务层、核心层和基础设施层四个层级。"
    "交互层包括CLI命令行界面和Web UI浏览器界面，负责接收用户输入并展示AI反馈。"
    "业务层包含Agent智能体系统、Session会话管理、Permission权限引擎等核心业务逻辑。"
    "核心层提供Tool工具系统、Skill知识注入、Memory记忆管理等可插拔能力模块。"
    "基础设施层包含Provider模型适配器、Database数据库、Python环境管理、License授权验证等底层支撑。"
    "系统整体架构如图3-1所示（此处为架构示意，实际代码中通过Effect-TS的Layer机制实现依赖注入和组合）。"
)

add_figure_caption('图3-1 Try系统总体架构图')

add_body(
    "数据流方面，用户通过CLI或Web UI输入消息后，系统首先通过会话管理层获取当前会话上下文"
    "和历史消息；然后由Agent执行引擎将上下文、工具列表、Skill知识组装成LLM调用请求；"
    "通过Provider适配层调用配置的LLM提供商API；AI响应以SSE流式方式返回，其中可能包含"
    "工具调用（Tool Call）指令；工具调用请求经过权限引擎审核后，通过Python子进程执行器"
    "调用对应的Python工具脚本；工具执行结果反馈给LLM进行下一轮推理，直到任务完成或"
    "达到最大迭代次数。"
)

add_section_title('3.2 Effect-TS函数式框架')

add_body(
    "Effect-TS是一个基于TypeScript的函数式编程框架，受Scala的ZIO库启发。"
    "它提供了强大的Effect类型系统，将副作用（如I/O操作、数据库查询、网络请求）"
    "编码为不可变的值描述，而非直接执行。Effect类型定义为Effect<Requirements, Error, Success>，"
    "明确表达了操作所需的环境依赖、可能产生的错误类型以及成功时的返回值类型。"
)

add_body(
    "Try系统全面采用Effect-TS的以下核心机制：（1）Context.Tag依赖注入："
    "通过定义AgentService、ToolRegistry、SessionService、SkillRegistry等服务标签，"
    "将组件间的依赖关系显式化，便于测试时注入Mock实现。"
    "（2）Layer分层组合：通过Effect.provide()和Layer.merge()机制将各服务模块组合为"
    "完整的运行时环境（AppLayer），实现声明式依赖管理。"
    "（3）Stream流式处理：利用Effect-Stream构建响应式数据管道，实现LLM响应的"
    "逐token流式推送和工具调用的实时反馈。"
    "（4）Exit类型安全错误处理：通过Either/Exit类型对所有可能失败的操作进行"
    "编译期检查，避免运行时未捕获异常。"
)

add_section_title('3.3 模块分层设计')

add_body(
    "系统各模块及其职责如下："
)

add_body(
    "（1）agent模块：包含Agent.ts（Agent服务核心逻辑）、executor.ts（迭代执行引擎，"
    "是系统的「心脏」）、registry.ts（Agent注册表）、types.ts（类型定义）、protocol.ts"
    "（Agent间通信协议）、builtin/（8个内置Agent实现）。Agent执行引擎实现了完整的状态机："
    "WAITING_FOR_INPUT → LLM_CALLING → TOOL_CALLING → TOOL_RUNNING → ERROR/HALTED/COMPLETED。"
)

add_body(
    "（2）tool模块：包含registry.ts（工具注册表）、loader.ts（Python工具加载器，"
    "通过Bun.spawn启动Python子进程，设置PYTHONIOENCODING=utf-8环境变量解决编码问题）、"
    "types.ts（类型定义）、confirmation.ts（操作确认管理）、builtin/（内置工具TypeScript入口）。"
    "tools/目录包含各工具的Python实现和TOOL.md元数据定义。"
)

add_body(
    "（3）session模块：包含session.ts（会话服务接口）、live.ts（SQLite持久化实现，"
    "创建sessions/messages表，sessions表新增workspace列支持工作目录隔离）。"
    "通过Effect-TS的Service模式分离接口定义与实现，支持环境变量覆盖配置。"
)

add_body(
    "（4）server模块：包含index.ts（Express风格HTTP服务器）、router.ts（路由注册）、"
    "middleware.ts（JWT认证/CORS/日志中间件）、websocket.ts（WebSocket实时通信）、"
    "handlers/（各API处理器：agent.ts处理Agent调用、chat.ts处理SSE流式对话、"
    "session.ts处理会话CRUD、workspace.ts处理工作目录API、auth.ts处理用户认证等）。"
)

add_body(
    "（5）provider模块：包含provider.ts（统一LLM调用接口，适配OpenAI/Anthropic/DeepSeek/"
    "Ollama/llama.cpp）、auth.ts（从auth.json加载API Key，支持命令行--auth-env从环境变量读取）。"
)

add_body(
    "（6）其他模块：skill模块负责加载和管理Markdown格式的知识注入文件；"
    "memory模块实现对话内容自动压缩（基于阈值和token计数）和嵌入检索（通过LLM生成嵌入向量）；"
    "permission模块实现基于规则的权限检查引擎；config模块管理try.json全局配置；"
    "infra模块提供基础设施（Python环境、日志、指标、迁移、License授权等）；"
    "web子项目使用Vite构建React前端，打包输出到dist/web/目录供服务器托管。"
)

new_page()

# ============================================================
# 第四章 Agent智能体系统
# ============================================================
add_chapter_title('第四章 Agent智能体系统')

add_section_title('4.1 Agent体系设计')

add_body(
    "Agent是本系统的核心抽象，代表一个具有特定角色、能力和工具集的AI智能体。"
    "每个Agent由以下属性定义：id（唯一标识符，如'builtin:builder'）、name（显示名称）、"
    "description（能力描述）、systemPrompt（系统提示词，定义Agent的行为模式和专业领域）、"
    "tools（可用的工具列表）、permissions（权限配置）、maxIterations（最大对话轮数）。"
)

add_body(
    "系统内置8种Agent角色：（1）Chat——通用对话Agent，拥有9个只读类工具（read_file、"
    "grep、glob、fetch_webpage等），适合代码探索、知识问答和任务规划。"
    "（2）Builder——全栈开发Agent，拥有12个完整工具（读写文件、Shell命令、代码搜索），"
    "是功能最全面的Agent，也是唯一可以委托子任务给其他Agent的角色。"
    "（3）Coder——专注代码编写与修改，根据指令生成或修改代码文件。"
    "（4）Reviewer——代码审查专家，分析代码中的潜在问题、安全隐患和最佳实践偏离。"
    "（5）Tester——测试工程师，运行测试套件并分析测试结果。"
    "（6）Refactor——代码重构专家，改善代码结构而不改变外部行为。"
    "（7）Researcher——信息检索专家，通过联网搜索和文档查阅获取外部知识。"
    "（8）Orchestrator——任务编排者，负责将复杂任务分解为子任务并协调多Agent执行。"
)

add_body(
    "Agent注册表（AgentRegistry）采用Effect-TS的Context.Tag机制，在应用启动时将"
    "所有Agent实现注册到DI容器中。CLI命令/agent list可列出所有可用Agent，"
    "/agent <id>可动态切换当前会话的Agent。Web UI通过AgentSelector下拉组件展示"
    "Agent列表并支持一键切换。Agent切换会在前端显示当前Agent的名称、描述和能力概览。"
)

add_section_title('4.2 执行引擎')

add_body(
    "执行引擎（src/agent/executor.ts）是系统的核心组件，负责驱动Agent与LLM之间的多轮交互。"
    "它实现了以下关键机制："
)

add_body(
    "（1）迭代循环：引擎以循环方式运行，每轮迭代执行LLM调用→解析响应→工具调用→结果反馈"
    "的标准流程。最大迭代次数由Agent配置决定（默认为20轮），防止无限循环消耗API配额。"
    "每轮迭代后将AI消息和工具调用结果追加到对话上下文中。"
)

add_body(
    "（2）状态管理：引擎维护一个执行状态机，包含以下阶段：WAITING_FOR_INPUT（等待用户输入）、"
    "LLM_CALLING（调用LLM中）、TOOL_CALLING（LLM返回工具调用指令）、TOOL_RUNNING（执行工具中）、"
    "RESPONDING（生成最终响应）、HALTED（异常终止）、ERROR（错误终止）、COMPLETED（正常完成）。"
    "每个阶段变更时通过Stream发送phase_change事件，CLI和Web UI据此更新界面状态指示。"
)

add_body(
    "（3）异常处理：当Agent达到最大迭代次数时，引擎首先通过setPhase('error', ...)向"
    "前端推送终止信息（包括已用轮数、最大轮数等详情），抛出MaxIterationsExceededError，"
    "再以Effect.fail()终止迭代。这一设计修正了早期版本中先fail()再推送导致前端无法接收"
    "终止通知的问题。当用户通过Web UI或CLI发送取消指令时，引擎通过AbortController中断"
    "当前的LLM调用和工具执行。"
)

add_body(
    "（4）Workspace集成：引擎在初始化时从会话数据中读取workspace路径，传递给工具执行器。"
    "如果会话未设置workspace，回退到defaultWorkspace()获取系统默认工作目录。"
    "这确保了Agent的文件操作始终在用户指定的工作目录内进行。"
)

add_section_title('4.3 Agent委托机制')

add_body(
    "Agent委托（Delegation）机制允许Builder和Orchestrator将复杂的子任务分配给更专业的Agent。"
    "委托通过内置的delegate工具实现，该工具接收目标Agent的ID和任务描述，创建"
    "子会话并运行目标Agent。委托流程如下：Builder收到用户的复杂需求（如「开发一个Web应用」）后，"
    "LLM判断任务可分解，调用delegate工具指定Coder编写核心代码→Tester运行测试→"
    "Reviewer审查代码质量→Refactor优化结构→Researcher查询第三方库文档。每个子Agent"
    "独立运行并返回结果摘要，Builder汇总各子任务结果后生成最终交付物。"
)

add_body(
    "委托机制的实现依赖于protocol.ts中定义的Agent间通信协议，通过标准化消息格式"
    "（包含task_id、parent_session、status、result等字段）实现子Agent的状态追踪和"
    "结果回传。这一机制使得用户只需与一个Agent交互即可完成复杂的多步骤任务。"
)

new_page()

# ============================================================
# 第五章 工具系统
# ============================================================
add_chapter_title('第五章 工具系统')

add_section_title('5.1 工具架构设计')

add_body(
    "工具系统是Agent与外部世界交互的桥梁，允许AI读写文件、搜索代码、执行命令、"
    "抓取网页等。工具的注册和调用采用分层架构：上层TypeScript入口负责工具发现、"
    "参数校验和权限检查；底层Python脚本负责实际执行。每个工具由TOOL.md元数据文件"
    "和main.py实现文件组成，放置于tools/builtin/<tool_name>/目录下。"
)

add_body(
    "TOOL.md文件定义了工具的元数据，包括：工具名称、功能描述、参数列表（名称、"
    "类型、是否必需、描述、枚举值等）、使用示例和注意事项。TOOL.md采用标准化的"
    "Markdown格式编写，便于人类阅读和程序解析。系统在启动时自动扫描tools/builtin/"
    "和tools/user/目录，动态加载所有工具。"
)

add_body(
    "工具注册通过ToolRegistry管理，每个工具被封装为ToolDefinition对象，包含"
    "name、description、parameters（JSON Schema格式，用于LLM的function calling）、"
    "execute函数。14个内置工具覆盖了软件开发全流程的需求。"
)

add_section_title('5.2 Python执行协议')

add_body(
    "本系统的一个重要设计决策是将所有工具的底层执行逻辑从TypeScript迁移到Python。"
    "这一决策基于以下考量：（1）Python在文件处理、文本操作、系统调用等任务上拥有"
    "更成熟的生态系统；（2）工具修改和调试更加便捷，用户无需重新编译TypeScript即可"
    "修改工具行为；（3）Python脚本可以独立于TypeScript主程序运行和测试。"
)

add_body(
    "跨进程通信采用JSON标准输入输出协议：TypeScript主程序通过Bun.spawn启动Python"
    "子进程，将工具参数序列化为JSON并通过stdin传递；Python脚本从sys.stdin读取JSON、"
    "执行操作、将结果以JSON格式写入sys.stdout；TypeScript读取stdout获取结果。"
    "这种协议设计简单、可调试（可在终端直接echo JSON到Python脚本测试），且完全避免了"
    "序列化库依赖。"
)

add_body(
    "Windows编码兼容性：在Windows平台上，Python子进程默认使用cp936（GBK）编码，"
    "而Bun/TypeScript使用UTF-8编码，导致中文路径下的文件操作出现乱码错误。"
    "本系统通过双层修复解决此问题：（1）在Bun.spawn的环境变量中设置PYTHONIOENCODING=utf-8"
    "和PYTHONUTF8=1；（2）在tools/builtin/_shared/encoding.py中创建编码兼容模块，"
    "强制设置sys.stdin/stdout/stderr的编码为UTF-8，并提供cp936兜底解码函数。"
    "所有Python工具脚本通过import _shared.encoding进行副作用导入，确保编码一致性。"
)

add_body(
    "Python虚拟环境管理：系统通过src/infra/python-env.ts实现Python解释器的智能查找。"
    "查找优先级为：项目根目录的.venv虚拟环境 > 系统PATH中的python > python3。"
    "在Windows上，python3.exe不存在是常见问题，系统通过优先查找python命令并验证版本"
    "的方式解决。工具目录下的requirements.txt会被自动检测并pip install安装依赖。"
)

add_section_title('5.3 内置工具实现')

add_body(
    "14个内置工具分为以下类别："
)

add_body(
    "文件操作工具（5个）：read_file——读取指定文件的全部或指定行范围内容，"
    "支持offset/limit分页读取大文件；write_file——创建新文件并写入内容（拒绝覆盖已有文件，"
    "保护用户数据）；edit_file——精确替换文件中的指定字符串，通过old_str/new_str精确匹配"
    "机制避免意外修改；file_exists——检查文件或目录是否存在；glob——使用minimatch模式"
    "匹配查找文件，支持**递归搜索和多种文件类型过滤。"
)

add_body(
    "代码搜索工具（2个）：grep——在项目文件中对指定的正则表达式模式进行文本搜索，"
    "支持includePattern文件过滤和ignoredFiles选项；read_command——执行系统命令并"
    "捕获其标准输出和标准错误输出。"
)

add_body(
    "执行工具（2个）：run_command——在shell中执行命令，支持工作目录指定和超时控制，"
    "执行前需通过权限引擎和确认机制审核；read_command——读取指定命令的输出（只读操作，安全性较高）。"
)

add_body(
    "网络工具（1个）：fetch_webpage——使用HTTP GET抓取网页内容，自动提取<title>"
    "和正文文本，支持超时控制和User-Agent设置。"
)

add_body(
    "知识工具（2个）：list_skills——列出当前系统中可用的所有Skill及其简要描述；"
    "get_skill——获取指定Skill的完整内容，使AI能够根据任务需要加载专业知识。"
)

add_body(
    "记忆工具（2个）：remember——将重要信息（如用户偏好、项目约定）存入持久化记忆库；"
    "recall——根据查询搜索历史记忆，返回相关度最高的记忆条目。"
)

add_body(
    "推理工具（1个）：think——允许AI在没有外部工具调用的情况下进行内部推理和规划，"
    "思考内容记录在对话上下文中但不直接展示给用户（除非在调试模式）。"
)

new_page()

# ============================================================
# 第六章 Skill与记忆系统
# ============================================================
add_chapter_title('第六章 Skill与记忆系统')

add_section_title('6.1 Skill知识注入')

add_body(
    "Skill系统是本项目设计的可扩展知识注入机制，允许用户在不需要修改核心代码的情况下"
    "为Agent提供领域专业知识。Skill的本质是结构化的Markdown文档，放置在skills/目录树下，"
    "包含引导Agent行为的指令、规范、模板和最佳实践。Agent在需要时通过get_skill工具"
    "动态获取Skill内容，将其纳入对话上下文。"
)

add_body(
    "Skill按来源分为三类：（1）builtin/——随项目发布的内置Skill，包括architecture-guide"
    "（项目架构指南，描述Effect-TS分层设计和依赖注入模式）、code-review（代码审查规范，"
    "包含审查清单和评分标准）、pr-template（PR提交模板，指导代码贡献流程）、"
    "tool-design（工具设计原则，规范新增工具的TOOL.md编写格式和Python脚本约定）。"
    "（2）user/——用户自定义Skill，如kraken（自建智能知识库服务的ingest和search工具使用指南）、"
    "write-ppt（PPT文档撰写指南）。"
    "（3）remote/——预留给远程Skill加载的场景。"
)

add_body(
    "Skill加载器（src/skill/loader.ts）在系统启动时扫描skills/目录，解析SKILL.md文件"
    "的YAML frontmatter和正文内容，建立索引（名称→路径→内容缓存）。"
    "Skill注册表（SkillRegistry）提供list()和get(name)两个接口，供list_skills和"
    "get_skill工具调用。Skill执行器（src/skill/executor.ts）支持将Skill内容与用户消息"
    "组合后进行语义搜索，找到最相关的Skill自动注入到Agent的system prompt中。"
    "这种机制使得Agent在不需要显式调用工具的情况下即可获得领域知识增强。"
)

add_section_title('6.2 记忆系统')

add_body(
    "记忆系统为Agent提供跨会话的上下文维护能力。系统实现了两种记忆模式："
)

add_body(
    "（1）短期记忆：每次对话的消息历史保存在会话对象的messages数组中，"
    "Agent执行引擎将最近N轮对话内容作为LLM调用的上下文。当对话长度超过模型"
    "上下文窗口时，记忆压缩器（src/memory/compressor.ts）自动触发压缩，"
    "调用LLM对较早的对话内容进行摘要提炼，以摘要替代原始消息，确保上下文窗口"
    "不溢出。压缩阈值由token计数模块动态计算。"
)

add_body(
    "（2）长期记忆：通过remember工具将重要信息写入持久化记忆库（SQLite memory表）。"
    "记忆记录包含标题、内容、标签、时间戳和嵌入向量。嵌入向量由LLM API生成"
    "（src/memory/embedding.ts），recall工具接收到查询后，通过嵌入向量的余弦相似度"
    "计算找到相关记忆，按相关度降序返回。系统还支持auto-memory机制，在对话结束后"
    "自动分析对话内容并提取值得记住的信息。"
)

add_body(
    "记忆表支持FTS5全文搜索索引，用户可通过关键词直接搜索历史记忆。"
    "记忆操作同样受权限系统管理，recall和remember工具需要相应的read/write权限。"
    "记忆数据与会话数据一起持久化在SQLite数据库中，重启不丢失。"
)

new_page()

# ============================================================
# 第七章 会话管理与安全体系
# ============================================================
add_chapter_title('第七章 会话管理与安全体系')

add_section_title('7.1 会话管理')

add_body(
    "会话（Session）是用户与Agent交互的基本单元。每个会话包含以下属性："
    "唯一ID（UUID v4）、标题（创建时自动生成，后续可由用户修改）、Agent配置ID、"
    "创建/更新时间戳、消息列表（messages数组，每条消息包含role、content、"
    "tool_calls等字段）、workspace工作目录路径。会话通过SQLite数据库持久化存储。"
)

add_body(
    "会话服务（SessionService）提供完整的CRUD操作：create（创建新会话）、get（获取单个会话）、"
    "list（列出所有会话，支持分页和排序）、update（更新会话属性，支持部分更新）、"
    "delete（删除会话及其关联消息）、addMessage（向会话追加消息）、"
    "updateWorkspace（更新会话的工作目录路径）。所有操作通过Effect-TS的Service模式"
    "定义，支持Live实现（SQLite）和Mock实现（内存），方便单元测试。"
)

add_body(
    "sessions表结构包含：id（TEXT PRIMARY KEY）、title（TEXT）、agent_id（TEXT）、"
    "workspace（TEXT NOT NULL DEFAULT ''）、messages（TEXT，JSON数组）、"
    "created_at、updated_at（ISO 8601时间戳）、user_id（TEXT，与用户认证关联）。"
    "workspace列通过ALTER TABLE ADD COLUMN迁移添加，兼容旧版本数据库。"
    "新建会话时workspace默认为系统默认工作路径。"
)

add_section_title('7.2 Workspace工作目录')

add_body(
    "Workspace是本系统引入的会话级工作目录隔离机制，确保Agent的所有文件操作在"
    "用户指定的目录范围内进行。Workspace的设计解决了两个核心问题："
    "（1）安全性——防止Agent访问项目源码、系统配置等敏感文件；"
    "（2）组织性——不同任务或项目的产出文件自动归类到各自的目录。"
)

add_body(
    "Workspace管理由src/infra/workspace.ts模块负责，提供三个核心函数："
    "defaultWorkspace()——返回项目根目录下的workspace/目录，如果不存在则自动创建；"
    "sanitizeWorkspace(raw)——对用户输入的路径进行安全校验，包括路径逃逸检测"
    "（防止../../../etc/passwd类型的攻击）和规范化（转换为绝对路径）；"
    "listWorkspaceSubdirs(base)——列出工作目录下的直接子目录，供前端选择器使用。"
)

add_body(
    "前端WorkspacePicker组件（src/web/src/components/WorkspacePicker.tsx）提供直观的"
    "工作目录管理界面：显示📁图标和缩写路径（取最后2段目录名）；点击进入编辑模式，"
    "输入框支持手动输入路径或从子目录下拉列表中选择；Enter确认/Esc取消；"
    "通过PUT /api/sessions/:id/workspace API持久化到数据库。Agent处理中（isProcessing）时"
    "自动禁用编辑，防止执行过程中的路径切换。"
)

add_section_title('7.3 权限与安全')

add_body(
    "权限系统是保障AI编程助手安全性的关键组件。系统实现了多层安全机制："
)

add_body(
    "（1）规则引擎（src/permission/rule-engine.ts）：核心权限判断基于Pattern-Rules模型。"
    "每个规则定义pattern（glob模式匹配文件路径或命令字符串）、allow（允许的操作列表："
    "read/write/edit/execute）、requireConfirm（是否需要用户手动确认）、"
    "description（规则说明）。规则按顺序匹配，命中第一个匹配的规则后停止。"
    "defaultAllow指定未匹配任何规则时的默认行为（通常仅允许read操作）。"
)

add_body(
    "（2）操作确认（src/tool/confirmation.ts）：当规则要求requireConfirm=true时，"
    "系统通过确认机制征求用户同意。在CLI模式下，终端显示操作详情并等待y/n确认；"
    "在Web UI模式下，浏览器弹出ConfirmDialog模态框展示操作详情。确认结果通过"
    "回调机制异步返回给执行引擎。"
)

add_body(
    "（3）Workspace隔离：执行文件操作前，verifyPathInWorkspace()函数验证目标路径"
    "是否位于会话的workspace目录内。对于编辑已有文件等需要访问workspace外部文件的"
    "操作（如read_file读取项目源码），需在权限规则中显式配置allow。"
)

add_body(
    "（4）认证授权：系统支持用户密码注册（POST /api/auth/register）和登录"
    "（POST /api/auth/login），密码使用bcrypt哈希存储。登录后返回JWT令牌，"
    "前端存储在localStorage中，后续请求通过Authorization: Bearer <token>头携带。"
    "JWT过期后需重新登录。中间件requireAuth()在受保护的API路由前验证令牌。"
)

new_page()

# ============================================================
# 第八章 前端与交互设计
# ============================================================
add_chapter_title('第八章 前端与交互设计')

add_section_title('8.1 React前端架构')

add_body(
    "系统前端从一个简单的Vanilla JS单页面应用全面重构为React 19应用，"
    "采用Vite 6构建工具和Chakra UI v3组件库。这一重构带来了显著的改进："
    "组件化架构提高了代码复用性和可维护性，声明式UI使得界面状态管理更加清晰，"
    "Chakra UI提供了开箱即用的无障碍支持和响应式设计。"
)

add_body(
    "前端核心组件架构如下：App.tsx是应用的根组件，管理全局状态（当前会话、"
    "workspace、Agent配置、认证信息等），协调子组件间的数据流。"
    "SessionSidebar——侧边栏组件，展示会话列表，支持创建/切换/删除会话，"
    "每个会话条目显示AI生成的自动标题和时间戳。"
    "ChatPanel——核心对话面板，渲染消息列表（ChatBubble组件），每条消息区分"
    "用户消息和AI回复，AI回复中支持Markdown渲染（react-markdown + remark-gfm）"
    "和代码语法高亮（react-syntax-highlighter / Prism）。"
)

add_body(
    "AgentSelector——Agent选择器，以下拉菜单展示8个内置Agent及其描述。"
    "WorkspacePicker——工作路径选择器，显示📁图标和路径缩写，点击可编辑，"
    "支持子目录下拉选择和Enter/Esc快捷键。"
    "SettingsDrawer——设置抽屉面板，支持模型配置（提供商/模型名/温度/最大令牌数切换）、"
    "权限规则编辑（JSON格式，实时语法校验）、环境变量查看。"
    "ConfirmDialog——操作确认弹窗，在敏感操作（如run_command执行危险命令）时显示。"
    "ToolSkillManager——工具与Skill管理面板，展示可用工具和Skill的详细信息。"
)

add_body(
    "数据通信方面，前端通过api.ts模块封装所有后端API调用（fetch封装，自动附加JWT头）。"
    "流式对话使用EventSource（SSE）连接，接收后端推送的chunks流（文本增量）和"
    "tool_calls事件（工具调用详情）。WebSocket连接通过useWebSocket Hook管理，"
    "用于实时状态推送。认证状态由AuthContext上下文全局管理。"
)

add_body(
    "构建产物通过Vite打包为静态文件（index.html + JS bundle + CSS），输出到"
    "dist/web/目录，由后端Bun HTTP服务器托管为静态资源。开发模式使用Vite dev server"
    "（bun run --cwd src/web dev）支持HMR热更新。"
)

add_section_title('8.2 CLI命令行界面')

add_body(
    "CLI模式通过Commander库实现命令行解析，提供三种启动方式："
    "bun start -- chat进入交互式REPL对话模式；bun start -- run <prompt>"
    "执行单次问答；bun start -- web启动Web服务器。REPL使用Node.js readline接口"
    "实现行编辑和命令历史，通过Chalk库实现彩色终端输出（用户消息/Agent响应/工具调用/"
    "错误信息使用不同颜色区分）。"
)

add_body(
    "CLI内置命令体系：/help查看帮助信息；/agent <id>切换Agent；/agent list列出"
    "所有Agent及描述；/model <name>切换模型；/file <path>添加文件到对话上下文；"
    "/clear清空当前对话；/exit退出程序。命令解析与Web API共享底层Agent执行逻辑，"
    "确保CLI和Web UI的行为一致性。"
)

new_page()

# ============================================================
# 第九章 部署与运维
# ============================================================
add_chapter_title('第九章 部署与运维')

add_section_title('9.1 Docker容器化')

add_body(
    "系统提供完整的Docker容器化部署方案。Dockerfile基于oven/bun:1官方镜像，"
    "额外安装Python 3、pip、venv和curl（用于健康检查）。构建流程为："
    "复制package.json和bun.lockb → bun install --frozen-lockfile安装依赖 → "
    "复制全部源代码、tools/、skills/、scripts/ → bun run build:web构建前端静态文件 → "
    "bun run build:compile编译TypeScript为独立二进制文件（dist/try-bin，注入RSA公钥）。"
    "运行时仅需执行./dist/try-bin web --host 0.0.0.0，无需Bun或Node.js环境。"
)

add_body(
    "docker-compose.yml提供了一键部署配置：build .构建镜像 → 映射3456端口 → "
    "通过环境变量注入API Keys（OPENAI_API_KEY等，也可通过auth.json配置） → "
    "挂载卷实现数据和配置持久化（./data数据库、./workspace工作目录、"
    "./model本地模型、./tools/user用户工具、./skills/user用户Skill）。"
    "健康检查通过curl -f http://localhost:3456/api/health每30秒检测一次。"
    "资源限制为最大2GB内存。"
)

add_section_title('9.2 构建与发布')

add_body(
    "构建脚本（scripts/build-backend.ts）实现了两阶段构建："
    "第一阶段——读取license_public.pem公钥文件，将PEM内容Base64编码后注入"
    "src/infra/license.ts的BUILD_PUBLIC_KEY_BASE64占位符；"
    "第二阶段——调用bun build编译TypeScript入口文件src/bin/try.ts，"
    "通过--target bun --compile标志编译为独立二进制文件（dist/try-bin.exe）。"
    "构建时排除node-llama-cpp和@node-llama-cpp/*（作为external依赖，在运行时动态加载）。"
    "编译完成后自动恢复license.ts原始内容（不污染git工作区）。"
)

add_body(
    "package.json定义了完整的npm scripts工具链：start启动开发模式、dev热重载开发、"
    "web启动Web服务器、dev:web启动前端开发服务器、build:web构建前端、"
    "build编译JS bundle、build:compile编译独立二进制、gen-keys生成RSA密钥对、"
    "sign-license签发许可证、test运行测试、typecheck TypeScript类型检查。"
)

new_page()

# ============================================================
# 第十章 测试与验证
# ============================================================
add_chapter_title('第十章 测试与验证')

add_section_title('10.1 功能测试')

add_body(
    "系统通过bun test运行自动化测试套件，覆盖以下核心模块："
)

add_body(
    "（1）Agent测试（testagent.test.ts）：验证Agent执行引擎的基本功能，包括单轮对话、"
    "多轮迭代的完整流程。测试覆盖正常完成、工具调用、最大迭代次数触发异常终止等场景，"
    "特别是验证了迭代终止时错误信息正确推送到前端的修复效果。"
)

add_body(
    "（2）工具测试（testool.test.ts和testool.ts）：验证每个内置工具的参数解析、"
    "Python子进程启动、JSON通信协议、执行结果返回的完整链路。重点测试了Windows"
    "中文路径下的文件操作正确性，确保编码兼容层（encoding.py + PYTHONIOENCODING=utf-8）生效。"
)

add_body(
    "（3）配置测试（testconfig.test.ts）：验证try.json配置文件的加载、解析、"
    "默认值合并、类型校验和环境变量覆盖逻辑。"
)

add_body(
    "（4）会话测试（testsession.test.ts）：验证会话CRUD操作、消息追加、"
    "workspace字段的读写持久化、FTS5搜索功能。包括旧数据库的workspace列迁移兼容性测试。"
)

add_body(
    "（5）权限测试（testpermission.test.ts）：验证规则引擎的匹配逻辑，覆盖精确匹配、"
    "glob模式匹配、defaultAllow回退、requireConfirm确认流程等场景。"
)

add_section_title('10.2 兼容性测试')

add_body(
    "多模型兼容性：系统在不同LLM提供商（OpenAI gpt-4o-mini、Anthropic claude-sonnet-4、"
    "DeepSeek deepseek-v4-flash、Ollama本地qwen2.5模型、llama.cpp加载GGUF模型）"
    "下均通过基础对话和工具调用测试，验证了Provider适配层的抽象正确性。"
)

add_body(
    "跨平台兼容性：系统在Windows（PowerShell 5.1）、macOS（bash/zsh）、"
    "Linux（bash）环境下分别测试，验证了路径处理、Python解释器查找、"
    "Shell命令执行等平台相关逻辑的正确性。特别是Windows的Python编码兼容层修复"
    "经过10/10全通过的端到端测试。"
)

add_body(
    "Python版本兼容：工具脚本在Python 3.10至3.12版本上均通过测试。"
    "Python虚拟环境的自动检测和使用（.venv优先 → python → python3回退）"
    "在Windows和POSIX系统上均正确工作。"
)

new_page()

# ============================================================
# 第十一章 总结与展望
# ============================================================
add_chapter_title('第十一章 总结与展望')

add_section_title('11.1 工作总结')

add_body(
    "本文设计并实现了一款开源、本地化、可扩展的AI编程助手系统——Try。"
    "主要完成了以下工作："
)

add_body(
    "第一，基于Effect-TS函数式框架构建了类型安全的系统架构，通过依赖注入、"
    "分层组合和流式处理等机制，确保了系统的可维护性和可测试性。"
    "系统包含Agent、Tool、Skill、Memory、Session、Permission、Provider、Server、"
    "CLI等十余个功能模块，各模块通过清晰的接口契约交互。"
)

add_body(
    "第二，设计了多角色Agent智能体体系，实现了Chat、Builder、Coder、Reviewer、"
    "Tester、Refactor、Researcher、Orchestrator共8种专用Agent。Agent执行引擎支持"
    "多轮迭代循环、状态管理和异常处理，并通过委托机制支持复杂任务的多Agent协作。"
    "修复了Agent达到最大迭代次数时错误信息不反馈前端的问题。"
)

add_body(
    "第三，完成了工具系统的Python重构，建立了JSON标准输入输出跨进程通信协议，"
    "解决了Windows平台中文路径编码兼容问题。14个内置工具覆盖文件操作、代码搜索、"
    "Shell执行、网页抓取、知识管理和记忆存储等完整开发工作流。"
)

add_body(
    "第四，实现了React 19 + Vite 6 + Chakra UI v3的现代化前端重构，"
    "提供会话管理、Agent切换、Workspace可视化选择、流式对话显示、"
    "工具调用可视化、权限确认等完整交互功能。同时维护了功能丰富的CLI交互界面。"
)

add_body(
    "第五，建立了完善的权限与安全体系，包括基于规则的细粒度权限引擎、"
    "敏感操作确认机制、Workspace工作目录隔离、用户认证与JWT授权、"
    "密码安全存储等安全措施。提供了Docker容器化部署方案和独立二进制编译能力。"
)

add_section_title('11.2 工作展望')

add_body(
    "尽管Try系统已经具备了AI编程助手的基本功能和良好的架构基础，但仍有以下方向"
    "值得进一步探索和改进："
)

add_body(
    "（1）智能化提升：引入RAG（检索增强生成）机制，将项目代码索引化并通过向量检索"
    "为Agent提供更深层的代码库理解能力。探索Agent自动学习和适应开发者编码风格的能力，"
    "提供个性化代码建议。实验更先进的Agent编排策略，如基于DAG的复杂任务自动拆解。"
)

add_body(
    "（2）多模态支持：扩展系统以支持图像理解（如UI截图→代码生成、架构图→代码骨架）。"
    "利用多模态LLM（如GPT-4o vision、Claude 3.5 Sonnet vision）增强Agent的感知能力。"
)

add_body(
    "（3）协作增强：实现多人协作模式，允许多个开发者共享同一Agent会话，"
    "支持实时代码审查和结对编程场景。开发VS Code深度集成插件，将Agent能力无缝嵌入"
    "到开发工作流中。"
)

add_body(
    "（4）性能优化：研究更高效的Agent执行策略，如并行工具调用、"
    "预测性工具预加载、增量对话压缩等。优化大型项目中的工具执行性能"
    "（如grep在百万行代码库中的搜索速度）。"
)

add_body(
    "（5）生态建设：建立工具和Skill的开放市场，供社区贡献和分享。"
    "支持更多LLM提供商（如Google Gemini、xAI Grok、国内的大模型平台等）。"
    "完善文档体系，提供在线演示环境和交互式教程，降低上手门槛。"
)

new_page()

# ============================================================
# 参考文献
# ============================================================
add_chapter_title('参考文献')
# 去掉 chapter title 的 first_line_indent
for p in doc.paragraphs:
    if '参考文献' in p.text and p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        p.paragraph_format.first_line_indent = Cm(0)

references = [
    "[1]  GitHub. GitHub Copilot: Your AI pair programmer[EB/OL]. https://github.com/features/copilot, 2024.",
    "[2]  Anysphere. Cursor: The AI-first Code Editor[EB/OL]. https://cursor.sh/, 2024.",
    "[3]  OpenCode Community. OpenCode: AI-powered coding assistant[CP/OL]. https://github.com/opencode-ai/opencode, 2024.",
    "[4]  Effect-TS. Effect: The missing standard library for TypeScript[CP/OL]. https://github.com/Effect-TS/effect, 2024.",
    "[5]  OpenAI. GPT-4 Technical Report[R]. arXiv preprint arXiv:2303.08774, 2023.",
    "[6]  Anthropic. The Claude Model Family[EB/OL]. https://docs.anthropic.com/, 2024.",
    "[7]  Bun. Bun: A fast all-in-one JavaScript runtime[EB/OL]. https://bun.sh/, 2024.",
    "[8]  Meta. React: A JavaScript library for building user interfaces[EB/OL]. https://react.dev/, 2024.",
    "[9]  Chakra UI. Chakra UI v3: Simple, modular and accessible component library[EB/OL]. https://chakra-ui.com/, 2024.",
    "[10] Vite. Vite: Next Generation Frontend Tooling[EB/OL]. https://vitejs.dev/, 2024.",
    "[11] Yao S, Zhao J, Yu D, et al. ReAct: Synergizing Reasoning and Acting in Language Models[C]. ICLR, 2023.",
    "[12] Wei J, Wang X, Schuurmans D, et al. Chain-of-Thought Prompting Elicits Reasoning in Large Language Models[C]. NeurIPS, 2022.",
    "[13] Ollama. Get up and running with large language models[EB/OL]. https://ollama.com/, 2024.",
    "[14] node-llama-cpp. Run LLMs locally with node.js[CP/OL]. https://github.com/withcatai/node-llama-cpp, 2024.",
    "[15] SQLite. SQLite: Small. Fast. Reliable. Choose any three[EB/OL]. https://www.sqlite.org/, 2024.",
    "[16] DeepSeek. DeepSeek-V4: A new generation of LLMs[EB/OL]. https://api.deepseek.com/, 2024.",
    "[17] continue.dev. Continue: The leading open-source AI code assistant[CP/OL]. https://github.com/continuedev/continue, 2024.",
    "[18] Paul Gauthier. Aider: AI pair programming in your terminal[CP/OL]. https://github.com/paul-gauthier/aider, 2024.",
    "[19] Docker Inc. Docker: Accelerated Container Application Development[EB/OL]. https://www.docker.com/, 2024.",
    "[20] 中国国家标准化管理委员会. GB/T 7714-2015 信息与文献 参考文献著录规则[S]. 北京: 中国标准出版社, 2015.",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(ref)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)  # 五号
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

new_page()

# ============================================================
# 致谢
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.space_before = Pt(36)
run = p.add_run('致    谢')
run.font.name = '黑体'
run.font.size = Pt(16)
run.font.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)

add_body(
    "在本次项目实训过程中，我得到了多方面的支持和帮助，谨在此表示诚挚的感谢。"
)

add_body(
    "首先，感谢指导老师在项目选题、架构设计和技术方案上给予的专业指导和宝贵建议。"
    "老师的悉心指导帮助我在面对复杂的系统设计问题时找到了正确的方向，使项目能够顺利完成。"
)

add_body(
    "感谢OpenCode开源项目提供的优秀架构参考，其基于Effect-TS的设计理念为本项目的"
    "技术选型和模块划分提供了重要的启发。感谢Effect-TS、Bun、React、Vite、Chakra UI等"
    "开源社区的贡献者，这些优秀的开源项目构成了本系统的技术基石，让我能够"
    "站在巨人的肩膀上进行创新。"
)

add_body(
    "感谢GitHub Copilot和Cursor等商业AI编程工具在实训过程中的辅助，"
    "它们的成功实践启发了本项目对AI编程助手痛点的思考，也让作者更加深刻地理解了"
    "构建开源替代方案的意义。"
)

add_body(
    "最后，感谢家人和朋友在实训期间的理解和支持。项目的开发过程充满挑战，"
    "从最初的架构设计到最终的代码调试，每一个阶段都凝聚着对技术的热爱和对完美的追求。"
    "这段实训经历不仅提升了我的工程实践能力，更让我深刻体会到开源精神和技术分享的价值。"
)

# ============================================================
# 保存
# ============================================================
output_path = r"d:\桌面\暑期实习\try\项目实习报告_Try_AI编程助手.docx"
doc.save(output_path)
print(f"✅ 报告已生成: {output_path}")
print(f"📄 总段落数: {len(doc.paragraphs)}")
# 统计字数
total_chars = sum(len(p.text) for p in doc.paragraphs)
print(f"📝 总字符数: {total_chars}")
